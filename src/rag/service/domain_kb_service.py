"""Domain pre-knowledge-base service.

Admin-uploaded files are stored as domain-scoped knowledge chunks. Semantic
completion uses these chunks as evidence before web search, without writing
back to A-side graph records.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from bs4 import BeautifulSoup
from sqlalchemy import text
from sqlalchemy.orm import Session

from src.rag.dependencies import ai_session_scope
from src.rag.service.embedding_service import search_text_embedding_chunks

PROJECT_ROOT = Path(__file__).resolve().parents[3]
UPLOAD_ROOT = PROJECT_ROOT / "data" / "domain_kb"
UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".html", ".htm", ".xml"}
MAX_TEXT_CHARS = 2_000_000
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180


def _json(value: Any) -> str:
    return json.dumps(value if value is not None else {}, ensure_ascii=False, sort_keys=True, default=str)


def _sid(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _hash_text(value: str) -> str:
    return hashlib.sha256((value or "").encode("utf-8")).hexdigest()


def _safe_filename(filename: str) -> str:
    name = Path(filename or "upload.txt").name.strip() or "upload.txt"
    return re.sub(r"[^0-9A-Za-z._\-\u4e00-\u9fff]+", "_", name)[:180]


def _decode_bytes(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"):
        try:
            return content.decode(encoding)
        except Exception:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> tuple[str, str | None]:
    try:
        from pypdf import PdfReader  # type: ignore
        import io

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages[:500]:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n\n".join(pages), None
    except Exception as exc:
        return "", f"pdf_extract_unavailable: {exc}"


def _extract_docx(content: bytes) -> tuple[str, str | None]:
    try:
        import docx  # type: ignore
        import io

        document = docx.Document(io.BytesIO(content))
        texts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    texts.append(" | ".join(cells))
        return "\n".join(texts), None
    except Exception as exc:
        return "", f"docx_extract_unavailable: {exc}"


def extract_text(filename: str, content: bytes) -> tuple[str, str | None]:
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        return _extract_docx(content)
    if suffix in TEXT_EXTENSIONS or not suffix:
        text_value = _decode_bytes(content)
        if suffix in {".html", ".htm", ".xml"}:
            soup = BeautifulSoup(text_value, "html.parser")
            for tag in soup(["script", "style", "noscript", "svg", "canvas"]):
                tag.decompose()
            text_value = soup.get_text("\n")
        return text_value[:MAX_TEXT_CHARS], None
    return "", f"unsupported_file_type: {suffix}"


def normalize_text(value: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in str(value or "").splitlines()]
    return "\n".join(line for line in lines if line)


def chunk_text(value: str, *, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    value = normalize_text(value)
    if not value:
        return []
    chunks: list[str] = []
    start = 0
    length = len(value)
    while start < length:
        end = min(length, start + chunk_size)
        slice_text = value[start:end].strip()
        if slice_text:
            chunks.append(slice_text)
        if end >= length:
            break
        start = max(0, end - overlap)
    return chunks


def ensure_domain_scenic(
    db: Session,
    *,
    source_scenic_id: str,
    source_scenic_pk: int | None = None,
    scenic_name: str | None = None,
) -> int:
    source_scenic_id = _sid(source_scenic_id)
    if not source_scenic_id:
        raise ValueError("source_scenic_id is required")
    row = db.execute(
        text(
            """
            insert into scenic_areas (source_scenic_id, source_scenic_pk, name, metadata, updated_at)
            values (:source_scenic_id, :source_scenic_pk, :name, cast(:metadata as jsonb), now())
            on conflict (source_scenic_id) do update set
                source_scenic_pk = coalesce(excluded.source_scenic_pk, scenic_areas.source_scenic_pk),
                name = coalesce(excluded.name, scenic_areas.name),
                metadata = scenic_areas.metadata || excluded.metadata,
                updated_at = now()
            returning id
            """
        ),
        {
            "source_scenic_id": source_scenic_id,
            "source_scenic_pk": source_scenic_pk,
            "name": scenic_name or source_scenic_id,
            "metadata": _json({"domain_kb_initialized": True}),
        },
    ).fetchone()
    return int(row[0])


def ingest_domain_kb_files(
    *,
    source_scenic_id: str,
    files: Iterable[tuple[str, bytes]],
    source_scenic_pk: int | None = None,
    scenic_name: str | None = None,
    submitted_by: str | None = None,
) -> dict[str, Any]:
    source_scenic_id = _sid(source_scenic_id)
    if not source_scenic_id:
        raise ValueError("source_scenic_id is required")
    now = datetime.now(timezone.utc).isoformat()
    saved_files: list[dict[str, Any]] = []
    total_chunks = 0
    with ai_session_scope() as db:
        scenic_id = ensure_domain_scenic(
            db,
            source_scenic_id=source_scenic_id,
            source_scenic_pk=source_scenic_pk,
            scenic_name=scenic_name,
        )
        for filename, content in files:
            safe_name = _safe_filename(filename)
            doc_id = uuid.uuid4().hex
            doc_dir = UPLOAD_ROOT / source_scenic_id / doc_id
            doc_dir.mkdir(parents=True, exist_ok=True)
            file_path = doc_dir / safe_name
            file_path.write_bytes(content)

            text_value, warning = extract_text(safe_name, content)
            chunks = chunk_text(text_value)
            inserted = 0
            for idx, chunk in enumerate(chunks, start=1):
                metadata = {
                    "doc_id": doc_id,
                    "filename": safe_name,
                    "chunk_index": idx,
                    "chunk_count": len(chunks),
                    "uploaded_at": now,
                    "submitted_by": submitted_by,
                    "storage_path": str(file_path),
                    "embedding_status": "pending",
                    "warning": warning,
                }
                db.execute(
                    text(
                        """
                        insert into knowledge_chunks (
                            scenic_id, source_scenic_id, source_type, source_id, source_node_id,
                            chunk_type, title, content, metadata, content_hash, sync_version,
                            source_table, source_pk, source_field, source_title, source_url,
                            evidence_text, created_at
                        ) values (
                            :scenic_id, :source_scenic_id, 'domain_kb', :source_id, '__domain__',
                            'domain_preknowledge', :title, :content, cast(:metadata as jsonb),
                            :content_hash, :sync_version, 'domain_kb_upload', :source_pk, 'file',
                            :source_title, :source_url, :evidence_text, now()
                        )
                        """
                    ),
                    {
                        "scenic_id": scenic_id,
                        "source_scenic_id": source_scenic_id,
                        "source_id": doc_id,
                        "title": f"{safe_name} #{idx}",
                        "content": chunk,
                        "metadata": _json(metadata),
                        "content_hash": _hash_text(f"{source_scenic_id}:{doc_id}:{idx}:{chunk}"),
                        "sync_version": "domain-kb-v1",
                        "source_pk": doc_id,
                        "source_title": safe_name,
                        "source_url": None,
                        "evidence_text": chunk[:500],
                    },
                )
                inserted += 1
            total_chunks += inserted
            file_row = {
                "doc_id": doc_id,
                "filename": safe_name,
                "size_bytes": len(content),
                "chunks": inserted,
                "status": "indexed" if inserted else "uploaded_no_text",
                "warning": warning,
                "embedding_status": "pending" if inserted else "skipped",
            }
            saved_files.append(file_row)

    return {
        "source_scenic_id": source_scenic_id,
        "files": saved_files,
        "total_files": len(saved_files),
        "total_chunks": total_chunks,
    }


def list_domain_kb_documents(source_scenic_id: str, *, limit: int = 100, offset: int = 0) -> dict[str, Any]:
    source_scenic_id = _sid(source_scenic_id)
    with ai_session_scope() as db:
        total = db.execute(
            text(
                """
                select count(distinct source_id) from knowledge_chunks
                where source_scenic_id=:source_scenic_id and source_type='domain_kb'
                """
            ),
            {"source_scenic_id": source_scenic_id},
        ).scalar() or 0
        rows = db.execute(
            text(
                """
                select source_id, min(source_title) as filename, count(*) as chunks,
                       min(created_at) as uploaded_at,
                       case
                           when count(*) filter (where metadata->>'embedding_status' = 'done') = count(*) then 'done'
                           when count(*) filter (where metadata->>'embedding_status' = 'failed') > 0 then 'failed'
                           else 'pending'
                       end as embedding_status,
                       max(metadata->>'embedding_model') as embedding_model,
                       max(metadata->>'embedding_error') as embedding_error,
                       max(metadata->>'warning') as warning
                from knowledge_chunks
                where source_scenic_id=:source_scenic_id and source_type='domain_kb'
                group by source_id
                order by min(created_at) desc
                limit :limit offset :offset
                """
            ),
            {"source_scenic_id": source_scenic_id, "limit": limit, "offset": offset},
        ).mappings().all()
    return {"items": [dict(row) for row in rows], "total": int(total)}


def _terms_from_query(query: str) -> list[str]:
    raw = re.split(r"[\s,，。；;、/|]+", str(query or ""))
    terms: list[str] = []
    for item in raw:
        cleaned = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "", item).strip().lower()
        if len(cleaned) >= 2 and cleaned not in terms:
            terms.append(cleaned)
    return terms[:16]


def search_domain_kb(source_scenic_id: str, query: str, *, limit: int = 5) -> list[dict[str, Any]]:
    source_scenic_id = _sid(source_scenic_id)
    terms = _terms_from_query(query)
    if not source_scenic_id or not terms:
        return []
    vector_results: list[dict[str, Any]] = []
    try:
        vector_results = search_text_embedding_chunks(source_scenic_id, query, limit=limit)
    except Exception:
        vector_results = []
    where_parts = []
    params: dict[str, Any] = {"source_scenic_id": source_scenic_id, "limit": limit}
    score_expr_parts = []
    for idx, term_value in enumerate(terms):
        key = f"term_{idx}"
        params[key] = f"%{term_value}%"
        where_parts.append(f"lower(content) like :{key}")
        score_expr_parts.append(f"case when lower(content) like :{key} then 1 else 0 end")
    where_sql = " or ".join(where_parts)
    score_sql = " + ".join(score_expr_parts) or "0"
    with ai_session_scope() as db:
        rows = db.execute(
            text(
                f"""
                select id, source_id, source_title, title, content, source_url,
                       evidence_text, metadata, ({score_sql}) as match_score
                from knowledge_chunks
                where source_scenic_id=:source_scenic_id
                  and source_type='domain_kb'
                  and ({where_sql})
                order by match_score desc, created_at desc
                limit :limit
                """
            ),
            params,
        ).mappings().all()
    keyword_results = []
    for row in rows:
        content = str(row.get("content") or "")
        keyword_results.append(
            {
                "title": row.get("title") or row.get("source_title") or "前置知识库",
                "content": content,
                "quote": row.get("evidence_text") or content[:500],
                "source": row.get("source_title") or "前置知识库",
                "source_type": "domain_kb_keyword",
                "source_url": None,
                "source_doc_id": row.get("source_id"),
                "chunk_id": row.get("id"),
                "page_no": int((row.get("metadata") or {}).get("chunk_index") or 0) or None,
                "score": float(row.get("match_score") or 0) + 1.5,
                "metadata": row.get("metadata") or {},
            }
        )

    merged: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in vector_results + keyword_results:
        key = str(item.get("source_url") or item.get("content") or "")
        if not key or key in seen:
            continue
        seen.add(key)
        merged.append(item)
    merged.sort(key=lambda x: float(x.get("score") or 0), reverse=True)
    return merged[:limit]



def delete_domain_kb_document(source_scenic_id: str, source_id: str) -> dict[str, Any]:
    """Delete one uploaded domain-KB document and all derived records.

    Removes text embeddings, queued/running job records, knowledge chunks, and
    the original stored file directory under data/domain_kb/{source_scenic_id}/{source_id}.
    """
    source_scenic_id = _sid(source_scenic_id)
    source_id = _sid(source_id)
    if not source_scenic_id or not source_id:
        raise ValueError("source_scenic_id and source_id are required")

    doc_dir = (UPLOAD_ROOT / source_scenic_id / source_id).resolve()
    allowed_root = (UPLOAD_ROOT / source_scenic_id).resolve()
    removed_files = False
    chunks_deleted = 0
    embeddings_deleted = 0
    jobs_deleted = 0

    with ai_session_scope() as db:
        chunk_rows = db.execute(
            text(
                """
                select id from knowledge_chunks
                where source_scenic_id=:source_scenic_id
                  and source_type='domain_kb'
                  and source_id=:source_id
                """
            ),
            {"source_scenic_id": source_scenic_id, "source_id": source_id},
        ).fetchall()
        chunk_ids = [int(row[0]) for row in chunk_rows]
        if chunk_ids:
            embeddings_deleted = int(
                db.execute(
                    text("delete from text_embeddings where chunk_id = any(:chunk_ids)"),
                    {"chunk_ids": chunk_ids},
                ).rowcount
                or 0
            )
        jobs_deleted = int(
            db.execute(
                text(
                    """
                    delete from domain_kb_embedding_jobs
                    where source_scenic_id=:source_scenic_id and source_id=:source_id
                    """
                ),
                {"source_scenic_id": source_scenic_id, "source_id": source_id},
            ).rowcount
            or 0
        )
        chunks_deleted = int(
            db.execute(
                text(
                    """
                    delete from knowledge_chunks
                    where source_scenic_id=:source_scenic_id
                      and source_type='domain_kb'
                      and source_id=:source_id
                    """
                ),
                {"source_scenic_id": source_scenic_id, "source_id": source_id},
            ).rowcount
            or 0
        )

    try:
        if str(doc_dir).startswith(str(allowed_root)) and doc_dir.exists() and doc_dir.is_dir():
            shutil.rmtree(doc_dir)
            removed_files = True
    except Exception as exc:
        return {
            "deleted": chunks_deleted > 0 or embeddings_deleted > 0 or jobs_deleted > 0,
            "source_scenic_id": source_scenic_id,
            "source_id": source_id,
            "chunks_deleted": chunks_deleted,
            "embeddings_deleted": embeddings_deleted,
            "jobs_deleted": jobs_deleted,
            "files_deleted": removed_files,
            "file_delete_error": str(exc),
        }

    return {
        "deleted": chunks_deleted > 0 or embeddings_deleted > 0 or jobs_deleted > 0 or removed_files,
        "source_scenic_id": source_scenic_id,
        "source_id": source_id,
        "chunks_deleted": chunks_deleted,
        "embeddings_deleted": embeddings_deleted,
        "jobs_deleted": jobs_deleted,
        "files_deleted": removed_files,
    }

def upsert_domain_shell(payload: dict[str, Any]) -> dict[str, Any]:
    source_scenic_id = _sid(payload.get("source_scenic_id")) or _sid(payload.get("code"))
    if not source_scenic_id:
        raise ValueError("source_scenic_id is required")
    source_scenic_pk = payload.get("source_scenic_pk")
    try:
        source_scenic_pk = int(source_scenic_pk) if source_scenic_pk not in (None, "") else None
    except Exception:
        source_scenic_pk = None
    name = _sid(payload.get("name")) or source_scenic_id
    metadata = dict(payload.get("metadata") or {})
    metadata.update({
        "source": "a_domain_create",
        "code": _sid(payload.get("code")),
        "description": _sid(payload.get("description")),
        "location": _sid(payload.get("location")),
    })
    with ai_session_scope() as db:
        scenic_id = ensure_domain_scenic(
            db,
            source_scenic_id=source_scenic_id,
            source_scenic_pk=source_scenic_pk,
            scenic_name=name,
        )
        db.execute(
            text(
                """
                update scenic_areas
                set metadata = coalesce(metadata, '{}'::jsonb) || cast(:metadata as jsonb),
                    updated_at = now()
                where id = :id
                """
            ),
            {"id": scenic_id, "metadata": _json(metadata)},
        )
        row = db.execute(
            text(
                """
                select id, source_scenic_id, source_scenic_pk, name, metadata
                from scenic_areas
                where id = :id
                """
            ),
            {"id": scenic_id},
        ).mappings().first()
    data = dict(row or {})
    data["created_or_updated"] = True
    return data

def _table_exists(db: Session, table_name: str) -> bool:
    return bool(db.execute(
        text("""
            select 1 from information_schema.tables
            where table_schema='public' and table_name=:table_name
            limit 1
        """),
        {"table_name": table_name},
    ).fetchone())


def _column_exists(db: Session, table_name: str, column_name: str) -> bool:
    return bool(db.execute(
        text("""
            select 1 from information_schema.columns
            where table_schema='public' and table_name=:table_name and column_name=:column_name
            limit 1
        """),
        {"table_name": table_name, "column_name": column_name},
    ).fetchone())


def _delete_if_exists(db: Session, table_name: str, where_sql: str, params: dict[str, Any]) -> int:
    if not _table_exists(db, table_name):
        return 0
    result = db.execute(text(f"delete from {table_name} where {where_sql}"), params)
    return int(result.rowcount or 0)


def delete_domain_all(source_scenic_id: str) -> dict[str, Any]:
    source_scenic_id = _sid(source_scenic_id)
    if not source_scenic_id:
        raise ValueError("source_scenic_id is required")

    counts: dict[str, int] = {}
    scenic_id: int | None = None
    domain_dir = (UPLOAD_ROOT / source_scenic_id).resolve()
    files_deleted = False
    file_delete_error = None

    with ai_session_scope() as db:
        row = db.execute(
            text("select id from scenic_areas where source_scenic_id=:sid order by id limit 1"),
            {"sid": source_scenic_id},
        ).fetchone()
        scenic_id = int(row[0]) if row else None

        if scenic_id is not None:
            chunk_rows = db.execute(
                text("select id from knowledge_chunks where scenic_id=:scenic_id or source_scenic_id=:sid"),
                {"scenic_id": scenic_id, "sid": source_scenic_id},
            ).fetchall() if _table_exists(db, "knowledge_chunks") else []
            chunk_ids = [int(item[0]) for item in chunk_rows]
            if chunk_ids and _table_exists(db, "text_embeddings") and _column_exists(db, "text_embeddings", "chunk_id"):
                counts["text_embeddings_by_chunk"] = int(db.execute(text("delete from text_embeddings where chunk_id = any(:chunk_ids)"), {"chunk_ids": chunk_ids}).rowcount or 0)

            for table in [
                "clip_image_embeddings",
                "image_embeddings",
                "text_embeddings",
                "node_property_claims",
                "semantic_edges",
                "semantic_nodes",
                "node_assets",
                "candidate_facts",
                "conflict_cases",
                "semantic_claim_candidates",
                "semantic_conflict_groups",
                "retrieval_hits",
                "retrieval_runs",
                "ingest_batches",
                "knowledge_chunks",
            ]:
                if _table_exists(db, table) and _column_exists(db, table, "scenic_id"):
                    counts[table] = counts.get(table, 0) + _delete_if_exists(db, table, "scenic_id=:scenic_id", {"scenic_id": scenic_id})

        for table in [
            "domain_kb_embedding_jobs",
            "node_property_claims",
            "semantic_edges",
            "semantic_nodes",
            "node_assets",
            "candidate_facts",
            "conflict_cases",
            "semantic_claim_candidates",
            "semantic_conflict_groups",
            "retrieval_runs",
            "ingest_batches",
            "knowledge_chunks",
        ]:
            if _table_exists(db, table) and _column_exists(db, table, "source_scenic_id"):
                counts[table] = counts.get(table, 0) + _delete_if_exists(db, table, "source_scenic_id=:sid", {"sid": source_scenic_id})

        if _table_exists(db, "sync_jobs"):
            job_rows = db.execute(text("select job_id from sync_jobs where source_scenic_id=:sid"), {"sid": source_scenic_id}).fetchall()
            job_ids = [str(item[0]) for item in job_rows]
            if job_ids and _table_exists(db, "sync_job_events"):
                counts["sync_job_events"] = int(db.execute(text("delete from sync_job_events where job_id = any(:job_ids)"), {"job_ids": job_ids}).rowcount or 0)
            counts["sync_jobs"] = int(db.execute(text("delete from sync_jobs where source_scenic_id=:sid"), {"sid": source_scenic_id}).rowcount or 0)

        if scenic_id is not None:
            counts["scenic_areas"] = int(db.execute(text("delete from scenic_areas where id=:scenic_id"), {"scenic_id": scenic_id}).rowcount or 0)
        else:
            counts["scenic_areas"] = int(db.execute(text("delete from scenic_areas where source_scenic_id=:sid"), {"sid": source_scenic_id}).rowcount or 0)

    try:
        allowed_root = UPLOAD_ROOT.resolve()
        if str(domain_dir).startswith(str(allowed_root)) and domain_dir.exists() and domain_dir.is_dir():
            shutil.rmtree(domain_dir)
            files_deleted = True
    except Exception as exc:
        file_delete_error = str(exc)

    return {
        "deleted": any(value > 0 for value in counts.values()) or files_deleted,
        "source_scenic_id": source_scenic_id,
        "scenic_id": scenic_id,
        "counts": counts,
        "files_deleted": files_deleted,
        "file_delete_error": file_delete_error,
    }

